# main.py
import asyncio
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from utils_helpers import extract_channel_from_text
from services_youtube_client import YouTubeAPIService
from orchestrators_master import MasterOrchestrator
from services_llm_service import get_llm_service
from services_embedding_service import get_embedding_service
from utils_logger import get_logger
from config_settings import settings

logger = get_logger("main")
app = FastAPI(title="YouTube Research Agent")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_event():
    logger.info("🚀 Initializing FULL MODE + HNSW...")
    # init LLM
    llm = get_llm_service({
        "groq_api_key": settings.groq_api_key,
        "groq_model": settings.groq_model
    })
    ok = await llm.test_connection()
    logger.info(f"LLM Test: {ok}")

    # init embedding (fastembed)
    get_embedding_service({"model": settings.sentence_transformers_model})
    logger.info("Embeddings loaded.")

@app.get("/")
def root():
    return {"status": "ok", "mode": "FULL_HNSW_METADATA"}

class QueryModel(BaseModel):
    query: str

# inside main.py (add imports at top)
# main.py or export_route.py
# Add imports at top of main.py
import json
from io import BytesIO
from fastapi.responses import HTMLResponse, PlainTextResponse, Response
from docx import Document

# ---------------------------------------------------------
# Helper: clean LLM output (remove ```json)
# ---------------------------------------------------------
def clean_llm_output(raw: str):
    if not raw:
        return raw
    raw = raw.replace("```json", "")
    raw = raw.replace("```", "")
    raw = raw.strip()
    return raw


# ---------------------------------------------------------
# Helper: Convert JSON report → nice readable text
# ---------------------------------------------------------
def json_to_text(report: dict) -> str:
    lines = []

    # ---------------------------------------------------------
    # 1. Automatic Channel Analysis (High-level overview)
    # ---------------------------------------------------------
    lines.append("AUTOMATIC CHANNEL ANALYSIS")
    lines.append("This report provides a structured overview of the channel based on metadata, semantic patterns, and video insights.")
    lines.append("")

    # Executive Summary
    if "executive_summary" in report:
        lines.append("EXECUTIVE SUMMARY")
        lines.append(report["executive_summary"])
        lines.append("")

    # ---------------------------------------------------------
    # 2. Content Themes
    # ---------------------------------------------------------
    if "themes" in report:
        lines.append("CONTENT THEMES")
        for t in report["themes"]:
            if isinstance(t, str):
                lines.append(f" • {t}")
            elif isinstance(t, dict):
                lines.append(f" • {t.get('name','')} ({t.get('engagement','')})")
        lines.append("")

    # ---------------------------------------------------------
    # 3. Engagement Metrics
    # ---------------------------------------------------------
    metrics = report.get("metrics", {})
    if metrics:
        lines.append("ENGAGEMENT METRICS")
        lines.append(f" • Subscribers: {metrics.get('subscriber_count','Not Available')}")
        lines.append(f" • Total Views: {metrics.get('total_views','Not Available')}")
        lines.append(f" • Total Videos: {metrics.get('total_videos','Not Available')}")
        lines.append(f" • Average Views: {metrics.get('average_views','Not Available')}")
        lines.append(f" • Engagement Rate: {metrics.get('engagement_rate','Not Available')}")
        lines.append("")

    # ---------------------------------------------------------
    # 4. Trending Topics
    # ---------------------------------------------------------
    if "trends" in report:
        lines.append("TRENDING TOPICS")
        for key, value in report["trends"].items():
            lines.append(f" • {key}: {value}")
        lines.append("")

    # ---------------------------------------------------------
    # 5. Top Recommendations
    # ---------------------------------------------------------
    if "top_recommendations" in report:
        lines.append("RECOMMENDATIONS")
        for r in report["top_recommendations"]:
            lines.append(f" • {r['title']} (Priority {r['priority']})")
        lines.append("")

    # ---------------------------------------------------------
    # 6. Actionable Tips
    # ---------------------------------------------------------
    if "short_actionable_tips" in report:
        lines.append("ACTIONABLE TIPS")
        for tip in report["short_actionable_tips"]:
            lines.append(f" • {tip}")
        lines.append("")

    # ---------------------------------------------------------
    # 7. Raw JSON (Optional for debugging)
    # ---------------------------------------------------------
    # lines.append("")
    # lines.append("RAW JSON (debug only):")
    # lines.append(json.dumps(report, indent=2))

    return "\n".join(lines)


@app.post("/export")
async def export(payload: dict, format: str = "html"):
    from io import BytesIO
    import json
    from fastapi.responses import HTMLResponse, PlainTextResponse, Response
    from docx import Document
    from docx.shared import RGBColor

    analysis = payload.get("analysis", {})
    raw_report = analysis.get("report", {})

    # Convert to dict if raw JSON text
    if isinstance(raw_report, str):
        raw_report = clean_llm_output(raw_report)
        try:
            raw_report = json.loads(raw_report)
        except:
            raw_report = {"executive_summary": raw_report}

    text_report = json_to_text(raw_report)

    channel = payload.get("channel", {})
    title = channel.get("snippet", {}).get("title", "YouTube Report")

    # ---------------------------------------------------------
    # HTML EXPORT (styled)
    # ---------------------------------------------------------
    if format == "html":
        html = f"""
        <html>
        <head>
            <meta charset="utf-8"/>
            <title>{title} Report</title>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
                h1 {{ color: #1E88E5; }}
                h2 {{ color: #1976D2; margin-top: 20px; }}
                pre {{ white-space: pre-wrap; }}
            </style>
        </head>
        <body>
            <h1>{title}</h1>
            <pre>{text_report}</pre>
        </body>
        </html>
        """
        return HTMLResponse(content=html, media_type="text/html")

    # ---------------------------------------------------------
    # TXT EXPORT
    # ---------------------------------------------------------
    if format == "txt":
        return PlainTextResponse(
            content=text_report,
            media_type="text/plain"
        )

    # ---------------------------------------------------------
    # DOCX EXPORT (with blue headings)
    # ---------------------------------------------------------
    if format == "docx":
        doc = Document()
        h = doc.add_heading(title, level=1)
        h.runs[0].font.color.rgb = RGBColor(0, 102, 204)

        for line in text_report.split("\n"):
            if line.endswith(":"):  # headings
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 102, 204)
            else:
                doc.add_paragraph(line)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return Response(
            content=buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={title}.docx"}
        )



@app.post("/query")
async def process_query(body: QueryModel):
    user_query = (body.query or "").strip()
    logger.info(f"User query: {user_query}")

    if not user_query:
        raise HTTPException(400, "query is required")

    channel_lookup = extract_channel_from_text(user_query)
    logger.info(f"Resolving channel: {channel_lookup}")

    try:
        yt = YouTubeAPIService()
        orchestrator = MasterOrchestrator()

        channel = yt.get_channel_by_name_or_handle(channel_lookup)
        if not channel:
            raise HTTPException(404, "Channel not found")

        channel_id = channel.get("id")
        videos = yt.get_channel_videos(channel_id, max_results=min(25, getattr(settings, "max_videos_per_channel", 25)))

        # build empty transcripts mapping (we do metadata-only embeddings)
        transcripts = {v.get("id"): "" for v in videos}

        result = await orchestrator.process(channel=channel, videos=videos, transcripts=transcripts)

        return {"channel": channel, "videos": videos, "analysis": result}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
