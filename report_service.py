# report_service.py
import io
import html
from typing import Any, Dict
from docx import Document
from docx.shared import Pt
from datetime import datetime
from utils_logger import get_logger

logger = get_logger("report_service")


def _safe_str(x):
    if x is None:
        return "Not Available"
    if isinstance(x, (dict, list)):
        try:
            import json
            return json.dumps(x, ensure_ascii=False, indent=2)
        except Exception:
            return str(x)
    return str(x)


# ---------------------------
# Plain text report generator
# ---------------------------
def build_text_report(channel: Dict[str, Any], analysis: Dict[str, Any]) -> bytes:
    lines = []
    title = channel.get("snippet", {}).get("title") or channel.get("title") or "Channel"
    lines.append(f"{title}")
    lines.append(f"Generated: {datetime.utcnow().isoformat()} UTC")
    lines.append("=" * 60)
    lines.append("\nCHANNEL METADATA\n")
    lines.append(f"Title: {title}")
    lines.append(f"Description: {_safe_str(channel.get('snippet', {}).get('description'))}")
    stats = channel.get("statistics", {}) or {}
    lines.append(f"Subscribers: {stats.get('subscriberCount', 'Not Available')}")
    lines.append(f"Total views: {stats.get('viewCount', 'Not Available')}")
    lines.append(f"Video count: {stats.get('videoCount', 'Not Available')}")
    lines.append("\nANALYSIS\n")
    report = analysis.get("report") if isinstance(analysis, dict) else analysis
    if isinstance(report, dict):
        # try to pick standard keys
        es = report.get("executive_summary") or report.get("summary") or report
        lines.append("Executive summary:\n")
        lines.append(_safe_str(es))
        lines.append("\nMetrics:\n")
        lines.append(_safe_str(report.get("metrics", "Not Available")))
        lines.append("\nThemes:\n")
        themes = report.get("themes") or []
        if themes:
            for t in themes:
                if isinstance(t, dict) and "name" in t:
                    lines.append(f"- {t.get('name')} (freq={t.get('frequency', '')})")
                else:
                    lines.append(f"- {t}")
        else:
            lines.append("Not Available")
        lines.append("\nRecommendations:\n")
        recs = report.get("recommendations") or []
        if recs:
            for r in recs:
                if isinstance(r, dict):
                    lines.append(f"- {r.get('title','')} — {r.get('description','')}")
                else:
                    lines.append(f"- {r}")
        else:
            lines.append("Not Available")
    else:
        # report is text
        lines.append(_safe_str(report))

    # semantic neighbors
    neighbors = analysis.get("seed_neighbors") or analysis.get("neighbors") or []
    if neighbors:
        lines.append("\nTop semantic neighbors (ids and similarity):")
        for n in neighbors:
            if isinstance(n, (list, tuple)):
                lines.append(f"- {n[0]} (dist={n[1]})")
            elif isinstance(n, dict):
                vid = n.get("video_id") or n.get("id") or "unknown"
                lines.append(f"- {vid} (dist={n.get('distance','')})")

    payload = "\n".join(lines)
    return payload.encode("utf-8")


# ---------------------------
# DOCX report generator
# ---------------------------
def build_docx_report(channel: Dict[str, Any], analysis: Dict[str, Any]) -> bytes:
    doc = Document()
    # Title
    title = channel.get("snippet", {}).get("title") or channel.get("title") or "Channel Report"
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()} UTC")

    # Channel metadata
    doc.add_heading("Channel metadata", level=2)
    desc = channel.get("snippet", {}).get("description") or ""
    p = doc.add_paragraph()
    p.add_run("Description: ").bold = True
    p.add_run(desc)

    stats = channel.get("statistics", {}) or {}
    p = doc.add_paragraph()
    p.add_run("Subscribers: ").bold = True
    p.add_run(str(stats.get("subscriberCount", "Not Available")))
    p.add_run("    ")
    p.add_run("Total views: ").bold = True
    p.add_run(str(stats.get("viewCount", "Not Available")))
    p.add_run("    ")
    p.add_run("Video count: ").bold = True
    p.add_run(str(stats.get("videoCount", "Not Available")))

    # Analysis
    doc.add_heading("Analysis", level=2)
    report = analysis.get("report") if isinstance(analysis, dict) else analysis
    if isinstance(report, dict):
        # executive summary
        doc.add_heading("Executive summary", level=3)
        es = report.get("executive_summary") or report.get("summary") or ""
        doc.add_paragraph(_safe_str(es))

        # metrics
        doc.add_heading("Metrics", level=3)
        metrics = report.get("metrics", {})
        for k, v in (metrics.items() if isinstance(metrics, dict) else []):
            doc.add_paragraph(f"{k}: {_safe_str(v)}", style="List Bullet")

        # themes
        doc.add_heading("Themes", level=3)
        themes = report.get("themes") or []
        for t in themes:
            if isinstance(t, dict):
                doc.add_paragraph(f"{t.get('name')} — freq: {t.get('frequency','')}, engagement: {t.get('engagement','')}", style="List Bullet")
            else:
                doc.add_paragraph(str(t), style="List Bullet")

        # recommendations
        doc.add_heading("Recommendations", level=3)
        recs = report.get("recommendations") or []
        for r in recs:
            if isinstance(r, dict):
                p = doc.add_paragraph(style="List Number")
                p.add_run(r.get("title", "Recommendation")).bold = True
                p.add_run("\n" + _safe_str(r.get("description", "")))
            else:
                doc.add_paragraph(str(r), style="List Number")
    else:
        doc.add_paragraph(_safe_str(report))

    # seed neighbors
    neighbors = analysis.get("seed_neighbors") or analysis.get("neighbors") or []
    if neighbors:
        doc.add_heading("Semantic neighbors", level=2)
        for n in neighbors:
            if isinstance(n, (list, tuple)):
                doc.add_paragraph(f"{n[0]} — distance: {n[1]}", style="List Bullet")
            elif isinstance(n, dict):
                vid = n.get("video_id") or n.get("id") or ""
                doc.add_paragraph(f"{vid} — distance: {n.get('distance','')}", style="List Bullet")

    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


# ---------------------------
# HTML report generator
# ---------------------------
def build_html_report(channel: Dict[str, Any], analysis: Dict[str, Any]) -> bytes:
    title = html.escape(channel.get("snippet", {}).get("title") or channel.get("title") or "Channel Report")
    desc = html.escape(channel.get("snippet", {}).get("description") or "")
    report = analysis.get("report") if isinstance(analysis, dict) else analysis
    now = datetime.utcnow().isoformat()

    # header
    html_parts = [
        "<!doctype html>",
        "<html><head><meta charset='utf-8'><title>Channel Report</title>",
        "<style>body{font-family:Arial,Helvetica,sans-serif;margin:24px}h1{color:#111}pre{background:#f6f6f6;padding:12px;border-radius:6px}</style>",
        "</head><body>"
    ]
    html_parts.append(f"<h1>{title}</h1>")
    html_parts.append(f"<p><em>Generated: {now} UTC</em></p>")
    html_parts.append("<h2>Channel metadata</h2>")
    html_parts.append(f"<p><strong>Description:</strong> {desc}</p>")
    stats = channel.get("statistics") or {}
    html_parts.append(f"<p><strong>Subscribers:</strong> {stats.get('subscriberCount', 'Not Available')} &nbsp;&nbsp; <strong>Total views:</strong> {stats.get('viewCount','Not Available')} &nbsp;&nbsp; <strong>Videos:</strong> {stats.get('videoCount','Not Available')}</p>")

    html_parts.append("<h2>Analysis</h2>")
    if isinstance(report, dict):
        html_parts.append("<h3>Executive summary</h3>")
        html_parts.append(f"<pre>{html.escape(_safe_str(report.get('executive_summary') or report.get('summary') or ''))}</pre>")

        html_parts.append("<h3>Metrics</h3>")
        html_parts.append(f"<pre>{html.escape(_safe_str(report.get('metrics', {})))}</pre>")

        html_parts.append("<h3>Themes</h3><ul>")
        for t in (report.get("themes") or []):
            if isinstance(t, dict):
                html_parts.append(f"<li>{html.escape(str(t.get('name')))} — freq: {t.get('frequency','')}</li>")
            else:
                html_parts.append(f"<li>{html.escape(str(t))}</li>")
        html_parts.append("</ul>")

        html_parts.append("<h3>Recommendations</h3><ul>")
        for r in (report.get("recommendations") or []):
            if isinstance(r, dict):
                html_parts.append(f"<li><strong>{html.escape(r.get('title',''))}</strong> — {html.escape(_safe_str(r.get('description','')))}</li>")
            else:
                html_parts.append(f"<li>{html.escape(str(r))}</li>")
        html_parts.append("</ul>")
    else:
        html_parts.append(f"<pre>{html.escape(_safe_str(report))}</pre>")

    # neighbors
    neighbors = analysis.get("seed_neighbors") or analysis.get("neighbors") or []
    if neighbors:
        html_parts.append("<h3>Semantic neighbors</h3><ul>")
        for n in neighbors:
            if isinstance(n, (list, tuple)):
                html_parts.append(f"<li>{html.escape(str(n[0]))} — {n[1]}</li>")
            elif isinstance(n, dict):
                vid = n.get("video_id") or n.get("id") or ""
                html_parts.append(f"<li>{html.escape(vid)} — {html.escape(_safe_str(n.get('distance','')))}</li>")
        html_parts.append("</ul>")

    html_parts.append("</body></html>")
    html_content = "\n".join(html_parts)
    return html_content.encode("utf-8")
